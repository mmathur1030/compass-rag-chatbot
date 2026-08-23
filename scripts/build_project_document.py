from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path("artifacts/Compass_Project_Documentation.docx")
INK = "000000"
MUTED = "555555"
BORDER = "DADCE0"


def set_run_font(run, size=11, bold=False, italic=False, color=INK):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    heading_tokens = {
        "Heading 1": (20, 20, 6, INK),
        "Heading 2": (16, 18, 6, INK),
        "Heading 3": (14, 16, 4, "434343"),
    }
    for name, (size, before, after, color) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    set_table_borders(table)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def style_table(table, alignments=None):
    repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.08
                if alignments:
                    paragraph.alignment = alignments[col_index]
                for run in paragraph.runs:
                    set_run_font(run, size=9.5, bold=row_index == 0)


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_body(doc, text):
    return doc.add_paragraph(text)


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_number(doc, text):
    return doc.add_paragraph(text, style="List Number")


def add_label_paragraph(doc, label, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(label)
    set_run_font(run, bold=True)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_page_break(doc):
    doc.add_page_break()


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    # Simple Google Docs-native opening block. Do not use Word's Title style.
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    title.paragraph_format.keep_with_next = True
    set_run_font(title.add_run("Compass"), size=26)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    set_run_font(subtitle.add_run("Enterprise Policy Q&A Bot"), size=16, color="434343")

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(24)
    set_run_font(meta.add_run("Week 2 Project Documentation  •  August 2026"), size=10, color=MUTED)

    lead = doc.add_paragraph()
    lead.paragraph_format.space_after = Pt(12)
    run = lead.add_run(
        "A measured retrieval-augmented generation application built with LangChain, "
        "LangGraph, Pinecone, Groq, and Streamlit."
    )
    set_run_font(run, size=12)

    add_label_paragraph(doc, "Purpose. ", "Help employees find grounded answers across HR, technical, and compliance policy documents.")
    add_label_paragraph(doc, "Core behavior. ", "Route, retrieve, rerank, answer with citations, and refuse unsupported questions.")
    add_label_paragraph(doc, "Evaluation. ", "Fifteen golden questions covering direct, exact-term, ambiguous, cross-document, and unanswerable cases.")

    add_heading(doc, "Executive summary", 1)
    add_body(
        doc,
        "Compass answers employee questions from three sample ACME policy manuals in a Streamlit "
        "chat interface. It separates the corpus into HR, Technical, and Compliance namespaces, "
        "then uses a LangGraph router to decide where each question should be searched."
    )
    add_body(
        doc,
        "Retrieval combines semantic Pinecone search with local BM25 keyword search. Reciprocal-rank "
        "fusion merges both rankings, a hosted reranker chooses the strongest evidence, and a relevance "
        "threshold sends weak matches to a deterministic refusal path. For supported questions, the "
        "Groq-hosted model receives numbered context blocks and must cite every factual claim."
    )
    add_body(
        doc,
        "The final golden run passed all 15 functional cases. Router accuracy, correct-source hit rate, "
        "citation validity, refusal accuracy, and automated faithfulness each reached 100%. The remaining "
        "performance gap is p95 latency: 12.78 seconds against an 8-second target."
    )

    add_heading(doc, "1. Problem and success criteria", 1)
    add_body(
        doc,
        "Employees often need answers that live in different policy manuals. A general-purpose language "
        "model may answer confidently without using the company's actual rules. Compass instead retrieves "
        "relevant passages first and asks the model to answer only from that evidence."
    )
    add_heading(doc, "Project one-liner", 2)
    add_body(
        doc,
        "Compass helps ACME employees answer HR, technical, and compliance questions from three sample "
        "enterprise manuals in a Streamlit chatbot, targeting at least 90% correct-source retrieval, 95% "
        "faithfulness, 100% valid citations and refusals, and an 8-second p95 response time."
    )
    add_heading(doc, "Success criteria", 2)
    for item in (
        "Route at least 90% of evaluation questions to the expected namespace or namespaces.",
        "Retrieve every expected source for at least 90% of answerable questions.",
        "Produce answers with at least 95% automated faithfulness to retrieved context.",
        "Use only valid numbered citations and correctly refuse every unanswerable gold question.",
        "Keep p95 end-to-end response latency at or below eight seconds.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "2. Dataset", 1)
    add_body(
        doc,
        "The demonstration corpus contains three English UTF-8 sample manuals. They are not live "
        "corporate systems of record. Together they produce 26 indexed chunks across three Pinecone "
        "namespaces."
    )
    dataset = doc.add_table(rows=1, cols=4)
    dataset.rows[0].cells[0].text = "Namespace"
    dataset.rows[0].cells[1].text = "Source"
    dataset.rows[0].cells[2].text = "Coverage"
    dataset.rows[0].cells[3].text = "Chunks"
    rows = [
        ("HR", "hr_policy.txt", "Leave, remote work, conduct, benefits, onboarding, discipline", "8"),
        ("Technical", "tech_docs.txt", "APIs, rate limits, architecture, deployment, webhooks, SLA", "9"),
        ("Compliance", "compliance_manual.txt", "Privacy, security, incidents, vendors, audits, recovery", "9"),
    ]
    for values in rows:
        cells = dataset.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    set_table_geometry(dataset, [1500, 2100, 4660, 1100])
    style_table(
        dataset,
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
    )

    add_heading(doc, "3. System architecture", 1)
    add_body(
        doc,
        "The application uses a stateful graph rather than one long chain. Each node has one job, and "
        "conditional edges make the refusal and answer paths explicit."
    )
    architecture = [
        ("Ingestion", "Load TXT, Markdown, and text-based PDF files; remove empty content; preserve source and page metadata."),
        ("Chunking", "Create 1,000-character chunks with 200-character overlap and deterministic SHA-256 IDs."),
        ("Embedding and storage", "Use Pinecone multilingual-e5-large embeddings with 1,024 dimensions and store vectors in three namespaces."),
        ("Routing", "Select HR, Technical, Compliance, or All using conversation history plus deterministic cross-domain detection."),
        ("Hybrid retrieval", "Collect up to eight dense and eight BM25 candidates per selected namespace, then fuse their rankings."),
        ("Reranking", "Use Pinecone bge-reranker-v2-m3 and keep the four strongest, diverse context chunks."),
        ("Decision", "Refuse when the top relevance score is below 0.05; otherwise continue to generation."),
        ("Generation", "Use Groq-hosted openai/gpt-oss-20b to answer only from numbered context blocks."),
        ("Verification", "Normalize citation syntax, validate citation ranges, and ask the model to repair invalid citations."),
        ("Interface and memory", "Show answers and sources in Streamlit; retain conversation history with an in-memory LangGraph checkpointer."),
    ]
    for label, text in architecture:
        add_label_paragraph(doc, f"{label}. ", text)

    add_heading(doc, "4. End-to-end RAG workflow", 1)
    workflow = (
        "Load supported source files and discard empty documents.",
        "Split each document into overlapping chunks and attach source metadata.",
        "Generate a stable ID and hosted embedding for every chunk.",
        "Upsert vectors into the corresponding HR, Technical, or Compliance namespace.",
        "Route the user's question to one namespace or all three.",
        "Run semantic dense search and exact-term BM25 search over the selected content.",
        "Fuse the rankings, remove duplicate candidates, and rerank the combined evidence.",
        "Return a safe refusal if the best evidence falls below the configured threshold.",
        "Otherwise send the four best numbered context blocks to the language model.",
        "Validate the answer's citations and repair them when necessary.",
        "Display the grounded answer, searched namespaces, and source passages in Streamlit.",
    )
    for item in workflow:
        add_number(doc, item)

    add_heading(doc, "5. Retrieval and routing design", 1)
    add_heading(doc, "Namespace routing", 2)
    add_body(
        doc,
        "Focused questions search one namespace, which limits irrelevant candidates and reduces work. "
        "Ambiguous or cross-document questions search all three. When multiple namespaces are selected, "
        "the final context first reserves the strongest available result from each namespace before filling "
        "the remaining context slots."
    )
    add_heading(doc, "Why hybrid retrieval", 2)
    add_body(
        doc,
        "Dense retrieval captures meaning even when the question and policy use different words. BM25 "
        "rewards exact terms such as error codes, API names, and policy phrases. Reciprocal-rank fusion "
        "combines these complementary rankings without assuming that their raw scores share a scale."
    )
    add_heading(doc, "Reranking and refusal", 2)
    add_body(
        doc,
        "The reranker scores the fused candidate pool against the original question. If the strongest "
        "relevance score is below 0.05, the graph does not call the answer generator. It returns a fixed "
        "message stating that the indexed documents do not contain sufficiently relevant information."
    )

    add_heading(doc, "6. Prompt and agent instructions", 1)
    add_heading(doc, "Router instructions", 2)
    add_body(
        doc,
        "The router defines HR, Technical, Compliance, and All. It uses message history to interpret "
        "follow-up questions and selects All when a query spans collections, is ambiguous, or cannot be "
        "assigned confidently. Explicit cross-domain terms provide an additional deterministic safeguard."
    )
    add_heading(doc, "Grounding instructions", 2)
    add_body(doc, "The answer-generation prompt applies five strict rules:")
    for item in (
        "Use the retrieved context as the only factual source.",
        "Cite every factual claim with one or more numbered sources.",
        "Do not add general knowledge, inferred requirements, or recommended best practices.",
        "Do not strengthen modal words—for example, changing “should” to “must.”",
        "Keep requirements from separate sources separate unless the context explicitly connects them.",
    ):
        add_bullet(doc, item)
    add_heading(doc, "Citation repair and evaluation judge", 2)
    add_body(
        doc,
        "The repair prompt permits only citation numbers present in the current context and forbids adding "
        "facts. A separate faithfulness prompt asks the evaluation model to score factual support only, not "
        "writing quality. Complete prompt text is version-controlled in graph.py and evaluation.py."
    )

    add_heading(doc, "7. Evaluation methodology", 1)
    add_body(
        doc,
        "The fixed gold set contains 15 questions: nine direct or exact-term questions, two cross-document "
        "questions, two ambiguous questions, and two questions the corpus cannot answer. Each example "
        "declares the expected namespace, source file, key answer facts, and answerability."
    )
    measures = [
        ("Router accuracy", "Whether all expected namespaces were selected."),
        ("Correct-source hit rate", "Whether every expected source file appeared in final retrieved context."),
        ("Keyword coverage", "Whether expected answer facts appeared after text normalization."),
        ("Refusal accuracy", "Whether answerable questions were answered and unanswerable questions were refused."),
        ("Citation validity", "Whether citations existed and every number mapped to a retrieved context block."),
        ("Faithfulness", "Whether an LLM judge found every factual claim supported by the retrieved context."),
        ("Latency", "Total graph execution time for each question, summarized by average and p95."),
    ]
    for label, text in measures:
        add_label_paragraph(doc, f"{label}. ", text)
    add_body(
        doc,
        "Important: faithfulness is an automated LLM-as-judge result produced by the same hosted model "
        "family used by the application. It is not an independent human audit."
    )

    add_heading(doc, "8. Final golden-test results", 1)
    add_body(doc, "The final run in evaluations/golden_run produced the following results.")
    metrics = doc.add_table(rows=1, cols=4)
    for index, value in enumerate(("Metric", "Target", "Result", "Status")):
        metrics.rows[0].cells[index].text = value
    metric_rows = [
        ("Router accuracy", "≥ 90%", "100.0%", "Met"),
        ("Correct-source hit rate", "≥ 90%", "100.0%", "Met"),
        ("Answer faithfulness", "≥ 95%", "100.0%", "Met"),
        ("Citation validity", "100%", "100.0%", "Met"),
        ("Correct refusal", "100%", "100.0%", "Met"),
        ("Answer keyword coverage", "Not set", "87.2%", "Observed"),
        ("Average response latency", "Not set", "5.04 s", "Observed"),
        ("p95 response latency", "≤ 8 s", "12.78 s", "Not met"),
    ]
    for values in metric_rows:
        cells = metrics.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    set_table_geometry(metrics, [3700, 1700, 1900, 2060])
    style_table(
        metrics,
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    for run in metrics.rows[-1].cells[3].paragraphs[0].runs:
        set_run_font(run, size=9.5, bold=True)

    add_heading(doc, "Results interpretation", 2)
    add_body(
        doc,
        "All 15 questions passed the current functional rubric. Dense-only and hybrid retrieval both "
        "achieved a 100% correct-source hit rate on this small corpus, so this evaluation does not show a "
        "measurable retrieval advantage for hybrid search. A larger, more confusable corpus with exact "
        "identifiers is needed for a meaningful comparison."
    )
    add_body(
        doc,
        "Tail latency is the only missed target. Questions routed to all namespaces were slowest because "
        "they perform more Pinecone searches and rerank a larger candidate pool. The 12.78-second p95 is "
        "also based on only 15 sequential requests and can vary with hosted-service and network conditions."
    )

    add_heading(doc, "9. Iterations and failure analysis", 1)
    iterations = (
        "Replaced the initial Chroma and OpenAI scaffold with Pinecone and Groq.",
        "Matched the Pinecone index dimension to the hosted embedding model.",
        "Separated a single default namespace into HR, Technical, and Compliance.",
        "Added LangGraph routing to avoid searching unrelated content for every question.",
        "Added BM25, reciprocal-rank fusion, and global reranking.",
        "Added a relevance threshold and deterministic refusal path.",
        "Added numbered context, citation validation, and conditional citation repair.",
        "Built a modern Streamlit chat interface with source display.",
        "Created a checkpointed 15-question evaluation harness and dense-versus-hybrid comparison.",
        "Tightened grounding instructions after one cross-domain answer added an unsupported security claim.",
    )
    for item in iterations:
        add_number(doc, item)
    add_heading(doc, "Observed failure and correction", 2)
    add_body(
        doc,
        "The cross-credential-security case initially stated that credentials were never logged or "
        "persisted insecurely, although the retrieved documents did not make that exact claim. The answer "
        "sounded reasonable but exceeded the evidence. The prompt was revised to prohibit inferred best "
        "practices, require support for every sentence, preserve modal wording, and keep requirements from "
        "different sources separate. The failed case was rerun and passed the faithfulness judge."
    )

    add_heading(doc, "10. Learnings", 1)
    learnings = (
        "Retrieval must be tested explicitly; a plausible answer does not prove that the correct evidence was found.",
        "Namespaces reduce noise for focused questions, but cross-domain questions need multi-namespace retrieval and deliberate source diversity.",
        "Dense and keyword retrieval solve different problems, although a small corpus can hide the benefit of combining them.",
        "A valid citation number does not guarantee that the cited passage supports the claim; citation validation and faithfulness evaluation are separate checks.",
        "The refusal threshold must be recalibrated whenever the corpus, embedding model, or reranker changes.",
        "Strict generation wording materially reduced unsupported inference.",
        "Tail latency—not average latency—is the remaining performance concern.",
    )
    for item in learnings:
        add_bullet(doc, item)

    add_heading(doc, "11. Limitations and next steps", 1)
    limitations = (
        "Move source documents to durable object storage and automate ingestion after updates.",
        "Track document versions and delete stale Pinecone chunks after source changes or removal.",
        "Replace the in-memory LangGraph checkpointer with durable conversation storage.",
        "Calibrate the refusal threshold on a larger labeled set and report precision and recall.",
        "Expand the corpus and gold set; add recall at k, ranking quality, and hybrid/reranking lift.",
        "Add human review or an independent judge model for faithfulness and citation correctness.",
        "Reduce p95 latency through parallel retrieval, caching, and profiling.",
        "Add authentication, authorization, audit logging, secret management, and document-level access controls before using real enterprise material.",
    )
    for item in limitations:
        add_bullet(doc, item)

    add_heading(doc, "12. Reproduction and project assets", 1)
    add_body(doc, "After configuring the environment and indexing the three namespaces, run:")
    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Inches(0.35)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(12)
    set_run_font(code.add_run("uv run pytest -q\nuv run rag-evaluate --output evaluations/golden_run"), size=10)
    add_body(
        doc,
        "The detailed report, machine-readable results, retrieval comparison, and checkpoint are stored "
        "in evaluations/golden_run/. The gold questions are in evaluations/questions.json. Core prompts "
        "and graph behavior are in src/rag_chatbot/graph.py; evaluation logic is in "
        "src/rag_chatbot/evaluation.py."
    )

    doc.core_properties.title = "Compass: Enterprise Policy Q&A Bot"
    doc.core_properties.subject = "Week 2 RAG Project Documentation"
    doc.core_properties.author = "Pranay Mathur"
    doc.core_properties.keywords = "RAG, LangChain, LangGraph, Pinecone, Groq, Streamlit"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
